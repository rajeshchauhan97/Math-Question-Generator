import os
from docx import Document
from docx.shared import Inches

def create_word_document(questions, output_path):
    """Create a Word document with questions and images."""
    doc = Document()
    
    # The loop will now run for each question, creating a separate document
    for i, question in enumerate(questions, start=1):
        # Add question metadata
        doc.add_heading(f"Question {i}", level=2)
        doc.add_paragraph(f"Title: {question.get('title', '')}")
        doc.add_paragraph(f"Description: {question.get('description', '')}")
        doc.add_paragraph(f"Difficulty: {question.get('difficulty', '')}")
        
        # Add question text
        doc.add_paragraph("\nQuestion:")
        doc.add_paragraph(question.get('question', ''))
        
        # Add image if exists
        image_path = question.get('image_path')
        if image_path and os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f"[Image load error: {str(e)}]")
        
        # Add options
        doc.add_paragraph("\nOptions:")
        for option in question.get('options', []):
            doc.add_paragraph(option, style='List Bullet') # Fix for UserWarning
        
        # Add correct answer
        doc.add_paragraph(f"\nCorrect Answer: {question.get('correct', '')}")
        
        # Add explanation
        doc.add_paragraph("\nExplanation:")
        doc.add_paragraph(question.get('explanation', ''))
        
        # Add curriculum info
        doc.add_paragraph("\nCurriculum Info:")
        doc.add_paragraph(f"Subject: {question.get('subject', '')}")
        doc.add_paragraph(f"Unit: {question.get('unit', '')}")
        doc.add_paragraph(f"Topic: {question.get('topic', '')}")
        doc.add_paragraph(f"Marks: {question.get('plusmarks', 1)}")
        
        # Since we're creating one document per question, no page break is needed
        # doc.add_page_break()
    
    doc.save(output_path)